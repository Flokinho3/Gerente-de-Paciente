"""
Aplicação Flask - Sistema de Gestão de Pacientes
Contém todas as rotas e lógica do Flask
"""
import json
import os
import io
import subprocess
import sys
import atexit
import threading
from typing import Dict
from flask import Flask, jsonify, render_template, request, send_file, make_response
from datetime import datetime
from database import Database

# Instância global do banco de dados (padrão)
_db_default = Database()

# Dicionário para armazenar instâncias de banco de dados por porta
_db_instances = {}

# Thread-local storage para armazenar a instância do banco de dados por thread
_thread_local = threading.local()

# Classe proxy para redirecionar chamadas para a instância correta do banco de dados
class DatabaseProxy:
    """Proxy que redireciona chamadas para a instância correta do banco de dados"""
    def __getattr__(self, name):
        db_instance = get_db()
        return getattr(db_instance, name)
    
    def __call__(self, *args, **kwargs):
        db_instance = get_db()
        return db_instance(*args, **kwargs)

# Instância proxy do banco de dados
db = DatabaseProxy()

# Carregar .env (pasta do .exe ou raiz do projeto)
from env_loader import load_env
load_env()

# Versão do sistema
VERSION = "1.0.3"
BUILD_DATE = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

app = Flask(__name__)

# Função helper para obter o banco de dados correto baseado na porta do request
def get_db():
    """Retorna a instância do banco de dados correta baseado na porta do request atual"""
    global db, _db_instances, _thread_local
    
    # Verificar se há uma instância configurada para esta thread
    if hasattr(_thread_local, 'db_instance'):
        return _thread_local.db_instance
    
    # Tentar obter a porta do servidor do request atual
    try:
        server_port = request.environ.get('SERVER_PORT')
        if server_port:
            port = int(server_port)
            if port in _db_instances:
                _thread_local.db_instance = _db_instances[port]
                return _db_instances[port]
    except:
        pass
    
    # Fallback: usar instância global padrão
    _thread_local.db_instance = _db_default
    return _db_default

# --- DISCOVERY=scan: estado em memória (peers + /register no líder) ---
_leader_lock = threading.Lock()
_discovery_peers = []   # [{"ip", "port", "ts"}, ...] atualizado pelo scan em leader.run_one_cycle
_registered_peers = []  # [{"ip", "port", "ts"}, ...] atualizado por POST /register (apenas no líder)


def _get_local_ip():
    import socket
    try:
        s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
        try:
            s.connect(('8.8.8.8', 80))
            return s.getsockname()[0]
        finally:
            s.close()
    except Exception:
        try:
            return socket.gethostbyname(socket.gethostname())
        except Exception:
            return '127.0.0.1'


def atualizar_discovery_peers(peers):
    """Atualiza a lista de peers (chamado por leader.run_one_cycle quando DISCOVERY=scan). peers: [{"ip", "port", "ts"}, ...]"""
    global _discovery_peers, _leader_lock
    with _leader_lock:
        _discovery_peers[:] = peers


# Injetar versão em todos os templates
@app.context_processor
def inject_version():
    return dict(version=VERSION, build_date=BUILD_DATE)

@app.route('/')
def home():
    return render_template('Home.html')

@app.route('/novo_paciente')
def novo_paciente():
    return render_template('novo_paciente.html')

@app.route('/pacientes')
def pacientes():
    return render_template('pacientes.html')

@app.route('/exportar')
def exportar():
    return render_template('exportar.html')

@app.route('/bd')
def bd():
    return render_template('bd.html')

@app.route('/conflitos')
def conflitos():
    return render_template('conflitos.html', version=VERSION)

@app.route('/agendamentos')
def agendamentos():
    return render_template('agendamentos.html')

@app.route('/aparencia')
def aparencia():
    return render_template('aparencia.html')

@app.route('/ajuda')
def ajuda():
    return render_template('ajuda.html')

@app.route('/api/version', methods=['GET'])
def get_version():
    """Retorna a versão atual do sistema"""
    return jsonify({
        'success': True,
        'version': VERSION,
        'build_date': BUILD_DATE
    })


@app.route('/health', methods=['GET'])
def health():
    """Health check para descoberta e liderança. Retorna ip, porta e status."""
    try:
        port = int(request.environ.get('SERVER_PORT') or os.getenv('PORT', 5000))
    except (ValueError, TypeError):
        port = int(os.getenv('PORT', 5000))
    ip = _get_local_ip()
    return jsonify({'ok': True, 'ip': ip, 'port': port, 'status': 'ok'})


def _ip_sort_key(ip_str, port=5000):
    """Chave de ordenação para menor IP (numérico). Usado por /register quando DISCOVERY=scan."""
    try:
        import ipaddress
        a = ipaddress.ip_address(ip_str)
        return (tuple(int(x) for x in str(a).split('.')), port)
    except Exception:
        return ((255, 255, 255, 255), port)


@app.route('/register', methods=['POST'])
def register():
    """Registro de peer no líder (usado por DISCOVERY=scan). Só o líder persiste; não-líder responde 200 sem efeito."""
    global _discovery_peers, _registered_peers, _leader_lock
    try:
        data = request.get_json() or {}
        ip = (data.get('ip') or '').strip()
        port = int(data.get('port', 5000))
    except (ValueError, TypeError):
        return jsonify({'ok': False, 'message': 'Payload inválido: {ip, port}'}), 400
    if not ip:
        return jsonify({'ok': False, 'message': 'ip obrigatório'}), 400
    my_ip = _get_local_ip()
    with _leader_lock:
        peers = list(_discovery_peers)
    if not peers:
        return jsonify({'ok': True, 'registered': False, 'message': 'Nenhum peer ainda'})
    min_peer = min(peers, key=lambda p: _ip_sort_key(p['ip'], p.get('port', 5000)))
    if min_peer['ip'] != my_ip:
        return jsonify({'ok': True, 'registered': False, 'message': 'Não sou o líder'})
    # Sou o líder: adicionar/atualizar em _registered_peers
    now = datetime.now().timestamp()
    with _leader_lock:
        for p in _registered_peers:
            if p.get('ip') == ip and p.get('port') == port:
                p['ts'] = now
                return jsonify({'ok': True, 'registered': True})
        _registered_peers.append({'ip': ip, 'port': port, 'ts': now})
    return jsonify({'ok': True, 'registered': True})


@app.route('/api/abrir_ajuda', methods=['GET'])
def abrir_ajuda():
    """Abre o arquivo de ajuda no Bloco de Notas do Windows"""
    try:
        caminho_ajuda = None
        
        # Detectar se está rodando como executável ou em desenvolvimento
        if getattr(sys, 'frozen', False):
            # Modo executável: procurar em vários locais possíveis
            # 1. No diretório temporário do PyInstaller (sys._MEIPASS)
            if hasattr(sys, '_MEIPASS'):
                caminho_temp = os.path.join(sys._MEIPASS, 'COMO_USAR.txt')
                if os.path.exists(caminho_temp):
                    caminho_ajuda = caminho_temp
            
            # 2. No diretório do executável
            if not caminho_ajuda:
                caminho_exe = os.path.join(os.path.dirname(sys.executable), 'COMO_USAR.txt')
                if os.path.exists(caminho_exe):
                    caminho_ajuda = caminho_exe
        else:
            # Modo desenvolvimento: usar o diretório do script
            caminho_ajuda = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'COMO_USAR.txt')
        
        # Verificar se o arquivo existe
        if not caminho_ajuda or not os.path.exists(caminho_ajuda):
            return jsonify({'success': False, 'message': 'Arquivo de ajuda não encontrado'}), 404

        # Abrir no Bloco de Notas do Windows
        subprocess.Popen(['notepad.exe', caminho_ajuda])
        
        return jsonify({'success': True, 'message': 'Arquivo de ajuda aberto no Bloco de Notas'})
    except Exception as e:
        return jsonify({'success': False, 'message': f'Erro ao abrir ajuda: {str(e)}'}), 500

@app.route('/api/salvar_paciente', methods=['POST'])
def salvar_paciente():
    """Salva os dados do paciente usando o banco de dados"""
    try:
        data = request.get_json()
        
        # Validar dados obrigatórios
        if not data or 'identificacao' not in data or 'avaliacao' not in data:
            return jsonify({'success': False, 'message': 'Dados inválidos'}), 400
        
        nome = data['identificacao'].get('nome_gestante', '').strip()
        if not nome:
            return jsonify({'success': False, 'message': 'Nome da gestante é obrigatório'}), 400
        
        # Salvar usando o banco de dados
        resultado = db.adicionar_paciente(data)
        
        return jsonify(resultado)
    except Exception as e:
        return jsonify({'success': False, 'message': f'Erro ao salvar: {str(e)}'}), 500

@app.route('/api/pacientes', methods=['GET'])
def listar_pacientes():
    """Lista todos os pacientes ou filtra por parâmetros"""
    try:
        filtro = {}
        
        # Filtros opcionais via query parameters
        nome = request.args.get('nome')
        unidade = request.args.get('unidade_saude')
        
        if nome:
            filtro['nome'] = nome
        if unidade:
            filtro['unidade_saude'] = unidade
        
        pacientes = db.buscar_pacientes(filtro if filtro else None)
        
        return jsonify({
            'success': True,
            'total': len(pacientes),
            'pacientes': pacientes
        })
    except Exception as e:
        import traceback
        error_details = traceback.format_exc()
        print(f"Erro ao buscar pacientes: {error_details}")
        return jsonify({
            'success': False, 
            'message': f'Erro ao buscar pacientes: {str(e)}',
            'error': str(e)
        }), 500

@app.route('/api/atualizar_paciente/<paciente_id>', methods=['PUT'])
def atualizar_paciente(paciente_id):
    """Atualiza os dados de um paciente existente"""
    try:
        data = request.get_json()
        
        # Validar dados obrigatórios
        if not data or 'identificacao' not in data or 'avaliacao' not in data:
            return jsonify({'success': False, 'message': 'Dados inválidos'}), 400
        
        # Atualizar usando o banco de dados
        resultado = db.atualizar_paciente(paciente_id, data)
        
        if resultado['success']:
            return jsonify(resultado)
        else:
            return jsonify(resultado), 404
        
    except Exception as e:
        return jsonify({'success': False, 'message': f'Erro ao atualizar: {str(e)}'}), 500

@app.route('/api/deletar_paciente/<paciente_id>', methods=['DELETE'])
def deletar_paciente(paciente_id):
    """Deleta um paciente do banco de dados"""
    try:
        resultado = db.deletar_paciente(paciente_id)
        
        if resultado['success']:
            return jsonify(resultado)
        else:
            return jsonify(resultado), 404
            
    except Exception as e:
        return jsonify({'success': False, 'message': f'Erro ao deletar: {str(e)}'}), 500

# Rotas de API para Agendamentos
@app.route('/api/agendamentos', methods=['GET'])
def listar_agendamentos():
    """Lista agendamentos com filtros opcionais"""
    try:
        paciente_id = request.args.get('paciente_id')
        data_inicio = request.args.get('data_inicio')
        data_fim = request.args.get('data_fim')
        status = request.args.get('status')
        
        agendamentos = db.listar_agendamentos(
            paciente_id=paciente_id,
            data_inicio=data_inicio,
            data_fim=data_fim,
            status=status
        )
        
        return jsonify({
            'success': True,
            'agendamentos': agendamentos
        })
    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Erro ao listar agendamentos: {str(e)}'
        }), 500

@app.route('/api/agendamentos', methods=['POST'])
def criar_agendamento():
    """Cria um novo agendamento"""
    try:
        data = request.get_json()
        
        # Validar dados obrigatórios
        paciente_id = data.get('paciente_id')
        data_consulta = data.get('data_consulta')
        hora_consulta = data.get('hora_consulta')
        
        if not paciente_id or not data_consulta or not hora_consulta:
            return jsonify({
                'success': False,
                'message': 'Dados obrigatórios faltando: paciente_id, data_consulta, hora_consulta'
            }), 400
        
        # Gerar ID único
        import uuid
        agendamento_id = str(uuid.uuid4())
        
        tipo_consulta = data.get('tipo_consulta')
        observacoes = data.get('observacoes')
        status = data.get('status', 'agendado')
        
        resultado = db.criar_agendamento(
            agendamento_id=agendamento_id,
            paciente_id=paciente_id,
            data_consulta=data_consulta,
            hora_consulta=hora_consulta,
            tipo_consulta=tipo_consulta,
            observacoes=observacoes,
            status=status
        )
        
        return jsonify(resultado)
    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Erro ao criar agendamento: {str(e)}'
        }), 500

@app.route('/api/agendamentos/<agendamento_id>', methods=['GET'])
def obter_agendamento(agendamento_id):
    """Obtém um agendamento específico"""
    try:
        agendamento = db.obter_agendamento(agendamento_id)
        if agendamento:
            return jsonify({
                'success': True,
                'agendamento': agendamento
            })
        else:
            return jsonify({
                'success': False,
                'message': 'Agendamento não encontrado'
            }), 404
    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Erro ao obter agendamento: {str(e)}'
        }), 500

@app.route('/api/agendamentos/<agendamento_id>', methods=['PUT'])
def atualizar_agendamento(agendamento_id):
    """Atualiza um agendamento existente"""
    try:
        data = request.get_json()
        
        resultado = db.atualizar_agendamento(
            agendamento_id=agendamento_id,
            data_consulta=data.get('data_consulta'),
            hora_consulta=data.get('hora_consulta'),
            tipo_consulta=data.get('tipo_consulta'),
            observacoes=data.get('observacoes'),
            status=data.get('status')
        )
        
        return jsonify(resultado)
    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Erro ao atualizar agendamento: {str(e)}'
        }), 500

@app.route('/api/agendamentos/<agendamento_id>', methods=['DELETE'])
def deletar_agendamento(agendamento_id):
    """Deleta um agendamento"""
    try:
        resultado = db.excluir_agendamento(agendamento_id)
        return jsonify(resultado)
    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Erro ao deletar agendamento: {str(e)}'
        }), 500

@app.route('/api/sync/discover', methods=['GET'])
def discover_servers():
    """Descobre outros servidores: Zeroconf (padrão) ou scan /24/SYNC_TARGETS (DISCOVERY=scan)"""
    try:
        import socket

        # Obter IP local e porta do servidor que atendeu o request
        s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
        try:
            s.connect(('8.8.8.8', 80))
            local_ip = s.getsockname()[0]
        except Exception:
            try:
                local_ip = socket.gethostbyname(socket.gethostname())
            except Exception:
                local_ip = '127.0.0.1'
        finally:
            s.close()
        try:
            sp = request.environ.get('SERVER_PORT')
            port = int(sp) if sp else int(os.getenv('PORT', 5000))
        except (ValueError, TypeError):
            port = int(os.getenv('PORT', 5000))

        flask_host = (os.getenv('FLASK_HOST') or '127.0.0.1').strip()
        host_warning = None
        if flask_host == '127.0.0.1':
            host_warning = (
                'Este servidor está escutando apenas em 127.0.0.1. Para outros PCs na rede '
                'serem encontrados e conectarem, defina FLASK_HOST=0.0.0.0 no .env e reinicie o Gerente. '
                'Em cada PC que rodar o Gerente, faça o mesmo. Verifique também se o Firewall do Windows '
                'permite o Gerente nas redes privadas.'
            )

        discovery = (os.getenv('DISCOVERY') or 'zeroconf').strip().lower()
        if discovery != 'scan':
            # Zeroconf: usar lista em memória do ServiceBrowser
            try:
                from inicio.rede.zeroconf_discovery import get_discovered_servers
                raw = get_discovered_servers()
                discovered_servers = [
                    s for s in raw
                    if not (str(s.get('ip')) == local_ip and int(s.get('port', 0)) == port)
                ]
            except Exception:
                discovered_servers = []
            resp = {'success': True, 'local_ip': local_ip, 'port': port, 'servers': discovered_servers}
            if host_warning:
                resp['host_warning'] = host_warning
            return jsonify(resp)

        # DISCOVERY=scan: varredura /24, SYNC_TARGETS ou SYNC_SCAN_CIDRS
        import ipaddress
        from concurrent.futures import ThreadPoolExecutor, as_completed

        discovered_servers = []
        targets = []
        targets_env = os.getenv('SYNC_TARGETS', '').strip()
        cidrs_env = os.getenv('SYNC_SCAN_CIDRS', '').strip()
        max_targets = int(os.getenv('SYNC_MAX_TARGETS', '1024'))

        def normalize_target(value):
            value = value.strip()
            if not value:
                return None
            try:
                ipaddress.ip_address(value)
                return value
            except ValueError:
                try:
                    return socket.gethostbyname(value)
                except Exception:
                    return None

        if targets_env:
            for raw in targets_env.replace(';', ',').split(','):
                ip = normalize_target(raw)
                if ip and ip != local_ip:
                    targets.append(ip)
        elif cidrs_env:
            for raw in cidrs_env.replace(';', ',').split(','):
                raw = raw.strip()
                if not raw:
                    continue
                try:
                    network = ipaddress.ip_network(raw, strict=False)
                except ValueError:
                    continue
                for host in network.hosts():
                    if len(targets) >= max_targets:
                        break
                    ip = str(host)
                    if ip != local_ip:
                        targets.append(ip)
        else:
            ip_parts = local_ip.split('.')
            base_ip = '.'.join(ip_parts[:-1])
            for i in range(1, 255):
                ip = f'{base_ip}.{i}'
                if ip != local_ip:
                    targets.append(ip)

        if len(targets) > max_targets:
            targets = targets[:max_targets]

        def check_server(ip):
            try:
                test_socket = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
                test_socket.settimeout(0.5)
                result = test_socket.connect_ex((ip, port))
                test_socket.close()
                if result == 0:
                    import urllib.request
                    try:
                        url = f'http://{ip}:{port}/api/version'
                        req = urllib.request.Request(url, timeout=1)
                        response = urllib.request.urlopen(req)
                        if response.status == 200:
                            try:
                                hostname = socket.gethostbyaddr(ip)[0]
                            except Exception:
                                hostname = ip
                            return {'ip': ip, 'port': port, 'hostname': hostname}
                    except Exception:
                        pass
            except Exception:
                pass

        workers = min(100, max(1, len(targets)))
        with ThreadPoolExecutor(max_workers=workers) as executor:
            futures = [executor.submit(check_server, ip) for ip in targets]
            for future in as_completed(futures):
                server = future.result()
                if server:
                    discovered_servers.append(server)

        resp = {'success': True, 'local_ip': local_ip, 'port': port, 'servers': discovered_servers}
        if host_warning:
            resp['host_warning'] = host_warning
        return jsonify(resp)
    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Erro ao descobrir servidores: {str(e)}'
        }), 500

@app.route('/api/sync/data', methods=['GET'])
def get_sync_data():
    """Retorna dados para sincronização (pacientes e agendamentos)"""
    try:
        from config import get_pc_id
        incluir_removidos = request.args.get('incluir_removidos', 'false').lower() == 'true'
        
        pacientes = db.buscar_pacientes(incluir_removidos=incluir_removidos)
        agendamentos = db.listar_agendamentos()
        
        return jsonify({
            'success': True,
            'pc_id': get_pc_id(),
            'pacientes': pacientes,
            'agendamentos': agendamentos,
            'timestamp': datetime.now().isoformat()
        })
    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Erro ao obter dados para sincronização: {str(e)}'
        }), 500

def _registros_identicos(reg1: Dict, reg2: Dict, campos_ignorados: set = None) -> bool:
    """Verifica se dois registros são idênticos (ignorando campos de sincronização)"""
    if campos_ignorados is None:
        campos_ignorados = {'ultima_modificacao', 'versao', 'pc_id'}
    
    campos_ignorados = campos_ignorados | {'id'}
    
    # Comparar todos os campos exceto os ignorados
    todos_campos = set(reg1.keys()) | set(reg2.keys())
    for campo in todos_campos:
        if campo in campos_ignorados:
            continue
        if reg1.get(campo) != reg2.get(campo):
            return False
    return True

def _marcar_conflito_paciente(paciente_id: str):
    """Marca um paciente como em conflito"""
    try:
        paciente = db.buscar_paciente(paciente_id)
        if paciente and paciente.get('status') != 'conflito':
            # Usar método do database para atualizar
            db_instance = get_db()
            cursor = db_instance.conn.cursor()
            cursor.execute("""
                UPDATE pacientes 
                SET status = 'conflito',
                    ultima_modificacao = ?,
                    versao = versao + 1
                WHERE id = ?
            """, (datetime.now().strftime('%Y-%m-%d %H:%M:%S'), paciente_id))
            db_instance.conn.commit()
    except Exception:
        pass

def _marcar_conflito_agendamento(agendamento_id: str):
    """Marca um agendamento como em conflito"""
    try:
        agendamento = db.obter_agendamento(agendamento_id)
        if agendamento and agendamento.get('status') != 'conflito':
            db_instance = get_db()
            cursor = db_instance.conn.cursor()
            cursor.execute("""
                UPDATE agendamentos 
                SET status = 'conflito',
                    ultima_modificacao = ?,
                    versao = versao + 1
                WHERE id = ?
            """, (datetime.now().strftime('%Y-%m-%d %H:%M:%S'), agendamento_id))
            db_instance.conn.commit()
    except Exception:
        pass

@app.route('/api/sync/merge', methods=['POST'])
def merge_sync_data():
    """Recebe dados de sincronização e faz merge inteligente com detecção de conflitos"""
    try:
        data = request.get_json()
        if not data:
            return jsonify({'success': False, 'message': 'Dados não fornecidos'}), 400
        
        pacientes_remotos = data.get('pacientes', [])
        agendamentos_remotos = data.get('agendamentos', [])
        pc_id_remoto = data.get('pc_id')  # PC_ID do servidor remoto
        
        # Obter pacientes locais e seus IDs (incluindo removidos para comparação)
        pacientes_locais = db.buscar_pacientes(incluir_removidos=True)
        pacientes_locais_dict = {p['id']: p for p in pacientes_locais}
        
        # Obter agendamentos locais (incluindo removidos para comparação)
        agendamentos_locais = db.listar_agendamentos()
        agendamentos_locais_dict = {a['id']: a for a in agendamentos_locais}
        
        # Estatísticas
        pacientes_adicionados = 0
        pacientes_atualizados = 0
        pacientes_conflito = 0
        agendamentos_adicionados = 0
        agendamentos_atualizados = 0
        agendamentos_conflito = 0
        
        # Sincronizar pacientes (registro por registro)
        for paciente_remoto in pacientes_remotos:
            paciente_id = paciente_remoto['id']
            paciente_local = pacientes_locais_dict.get(paciente_id)
            
            if not paciente_local:
                # 1. Registro não existe localmente → Adicionar
                try:
                    db.inserir_registro(
                        paciente_id,
                        paciente_remoto,
                        arquivo_origem='sync',
                        data_salvamento=paciente_remoto.get('data_salvamento'),
                        pc_id=paciente_remoto.get('pc_id', pc_id_remoto),
                        ultima_modificacao=paciente_remoto.get('ultima_modificacao'),
                        versao=paciente_remoto.get('versao', 1),
                        status=paciente_remoto.get('status', 'ativo')
                    )
                    pacientes_adicionados += 1
                except Exception:
                    pass
            else:
                # 2. Registro existe localmente
                status_local = paciente_local.get('status', 'ativo')
                status_remoto = paciente_remoto.get('status', 'ativo')
                
                # 2a. Verificar se um foi removido e outro editado → CONFLITO
                if (status_local == 'removido' and status_remoto != 'removido') or \
                   (status_local != 'removido' and status_remoto == 'removido'):
                    _marcar_conflito_paciente(paciente_id)
                    pacientes_conflito += 1
                    continue
                
                # 2b. Se ambos removidos → ignorar
                if status_local == 'removido' and status_remoto == 'removido':
                    continue
                
                # 2c. Verificar se são idênticos → Ignorar
                if _registros_identicos(paciente_local, paciente_remoto):
                    continue
                
                # 2d. São diferentes → Comparar ultima_modificacao
                ultima_mod_local = paciente_local.get('ultima_modificacao') or paciente_local.get('data_salvamento', '')
                ultima_mod_remota = paciente_remoto.get('ultima_modificacao') or paciente_remoto.get('data_salvamento', '')
                
                if ultima_mod_remota > ultima_mod_local:
                    # Remoto mais recente → Atualizar local
                    try:
                        paciente_remoto['pc_id'] = paciente_remoto.get('pc_id', pc_id_remoto)
                        db.atualizar_paciente(paciente_id, paciente_remoto)
                        pacientes_atualizados += 1
                    except Exception:
                        pass
                elif ultima_mod_local > ultima_mod_remota:
                    # Local mais recente → Manter local
                    pass
                else:
                    # Iguais (ou ambos vazios) → CONFLITO
                    # Desempate por pc_id se possível
                    pc_id_local = paciente_local.get('pc_id', '')
                    pc_id_remota = paciente_remoto.get('pc_id', pc_id_remoto or '')
                    
                    if pc_id_local != pc_id_remota:
                        # Se pc_ids diferentes, ainda pode ser conflito se dados diferentes
                        _marcar_conflito_paciente(paciente_id)
                        pacientes_conflito += 1
                    else:
                        # Mesmo pc_id, mesmo timestamp, mas dados diferentes → CONFLITO suspeito
                        _marcar_conflito_paciente(paciente_id)
                        pacientes_conflito += 1
        
        # Sincronizar agendamentos (registro por registro)
        for agendamento_remoto in agendamentos_remotos:
            agendamento_id = agendamento_remoto['id']
            agendamento_local = agendamentos_locais_dict.get(agendamento_id)
            
            if not agendamento_local:
                # 1. Registro não existe localmente → Adicionar
                try:
                    db.criar_agendamento(
                        agendamento_id,
                        agendamento_remoto['paciente_id'],
                        agendamento_remoto['data_consulta'],
                        agendamento_remoto['hora_consulta'],
                        agendamento_remoto.get('tipo_consulta'),
                        agendamento_remoto.get('observacoes'),
                        agendamento_remoto.get('status', 'agendado'),
                        agendamento_remoto.get('data_criacao'),
                        agendamento_remoto.get('data_atualizacao'),
                        pc_id=agendamento_remoto.get('pc_id', pc_id_remoto),
                        ultima_modificacao=agendamento_remoto.get('ultima_modificacao'),
                        versao=agendamento_remoto.get('versao', 1)
                    )
                    agendamentos_adicionados += 1
                except Exception:
                    pass
            else:
                # 2. Registro existe localmente
                status_local = agendamento_local.get('status', 'agendado')
                status_remoto = agendamento_remoto.get('status', 'agendado')
                
                # 2a. Verificar se um foi removido e outro editado → CONFLITO
                if (status_local == 'removido' and status_remoto != 'removido') or \
                   (status_local != 'removido' and status_remoto == 'removido'):
                    _marcar_conflito_agendamento(agendamento_id)
                    agendamentos_conflito += 1
                    continue
                
                # 2b. Se ambos removidos → ignorar
                if status_local == 'removido' and status_remoto == 'removido':
                    continue
                
                # 2c. Verificar se são idênticos → Ignorar
                if _registros_identicos(agendamento_local, agendamento_remoto, {'nome_gestante', 'unidade_saude'}):
                    continue
                
                # 2d. São diferentes → Comparar ultima_modificacao
                ultima_mod_local = agendamento_local.get('ultima_modificacao') or agendamento_local.get('data_atualizacao', '')
                ultima_mod_remota = agendamento_remoto.get('ultima_modificacao') or agendamento_remoto.get('data_atualizacao', '')
                
                if ultima_mod_remota > ultima_mod_local:
                    # Remoto mais recente → Atualizar local
                    try:
                        db.atualizar_agendamento(
                            agendamento_id,
                            agendamento_remoto['paciente_id'],
                            agendamento_remoto['data_consulta'],
                            agendamento_remoto['hora_consulta'],
                            agendamento_remoto.get('tipo_consulta'),
                            agendamento_remoto.get('observacoes'),
                            agendamento_remoto.get('status', 'agendado'),
                            pc_id=agendamento_remoto.get('pc_id', pc_id_remoto),
                            ultima_modificacao=agendamento_remoto.get('ultima_modificacao')
                        )
                        agendamentos_atualizados += 1
                    except Exception:
                        pass
                elif ultima_mod_local > ultima_mod_remota:
                    # Local mais recente → Manter local
                    pass
                else:
                    # Iguais (ou ambos vazios) → CONFLITO
                    _marcar_conflito_agendamento(agendamento_id)
                    agendamentos_conflito += 1
        
        return jsonify({
            'success': True,
            'message': 'Sincronização concluída',
            'stats': {
                'pacientes_adicionados': pacientes_adicionados,
                'pacientes_atualizados': pacientes_atualizados,
                'pacientes_conflito': pacientes_conflito,
                'agendamentos_adicionados': agendamentos_adicionados,
                'agendamentos_atualizados': agendamentos_atualizados,
                'agendamentos_conflito': agendamentos_conflito
            }
        })
    except Exception as e:
        import traceback
        return jsonify({
            'success': False,
            'message': f'Erro ao sincronizar: {str(e)}',
            'traceback': traceback.format_exc()
        }), 500

@app.route('/api/sync/conflitos', methods=['GET'])
def listar_conflitos():
    """Lista todos os conflitos pendentes"""
    try:
        resultado = db.listar_conflitos()
        return jsonify(resultado)
    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Erro ao listar conflitos: {str(e)}'
        }), 500

@app.route('/api/sync/conflitos/resolver', methods=['POST'])
def resolver_conflito():
    """Resolve um conflito específico"""
    try:
        data = request.get_json()
        if not data:
            return jsonify({'success': False, 'message': 'Dados não fornecidos'}), 400
        
        registro_id = data.get('registro_id')
        tipo = data.get('tipo')  # 'paciente' ou 'agendamento'
        acao = data.get('acao')  # 'manter_local', 'aceitar_remoto', 'mesclar'
        dados_remotos = data.get('dados_remotos')
        
        if not registro_id or not tipo or not acao:
            return jsonify({'success': False, 'message': 'Dados incompletos'}), 400
        
        resultado = db.resolver_conflito(registro_id, tipo, acao, dados_remotos)
        return jsonify(resultado)
    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Erro ao resolver conflito: {str(e)}'
        }), 500

@app.route('/api/sync/remover_pacientes', methods=['POST'])
def remover_pacientes_confirmados():
    """Remove pacientes após confirmação do usuário"""
    try:
        data = request.get_json()
        if not data:
            return jsonify({'success': False, 'message': 'Dados não fornecidos'}), 400
        
        paciente_ids = data.get('paciente_ids', [])
        if not paciente_ids:
            return jsonify({'success': False, 'message': 'Nenhum ID fornecido'}), 400
        
        resultado = db.remover_pacientes(paciente_ids)
        
        return jsonify(resultado)
    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Erro ao remover pacientes: {str(e)}'
        }), 500

@app.route('/api/backup/criar', methods=['GET'])
def criar_backup():
    """Cria um backup completo do banco de dados"""
    try:
        resultado = db.criar_backup()
        return jsonify(resultado)
    except Exception as e:
        return jsonify({'success': False, 'message': f'Erro ao criar backup: {str(e)}'}), 500

@app.route('/api/backup/download', methods=['GET'])
def download_backup():
    """Permite baixar o último backup como arquivo JSON"""
    try:
        resultado = db.criar_backup()
        buffer = io.BytesIO(json.dumps(resultado['backup'], ensure_ascii=False, indent=2).encode('utf-8'))
        buffer.seek(0)
        filename = f"backup_pacientes_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
        response = make_response(send_file(
            buffer,
            mimetype='application/json',
            as_attachment=True,
            download_name=filename
        ))
        return response
    except Exception as e:
        return jsonify({'success': False, 'message': f'Erro ao gerar arquivo de backup: {str(e)}'}), 500

@app.route('/api/backup/restaurar', methods=['POST'])
def restaurar_backup():
    """Restaura o banco de dados a partir de um backup"""
    try:
        data = request.get_json()
        if not data:
            return jsonify({'success': False, 'message': 'Dados do backup não fornecidos'}), 400

        backup = data.get('backup') if isinstance(data, dict) else data
        if not isinstance(backup, list):
            return jsonify({'success': False, 'message': 'Backup deve ser uma lista de registros'}), 400

        resultado = db.restaurar_backup(backup)
        
        if resultado['success']:
            return jsonify(resultado)
        else:
            return jsonify(resultado), 400
            
    except Exception as e:
        return jsonify({'success': False, 'message': f'Erro ao restaurar backup: {str(e)}'}), 500

@app.route('/api/backup/validar', methods=['POST'])
def validar_backup():
    """Valida a estrutura do backup antes de restaurar"""
    try:
        data = request.get_json()
        if not data:
            return jsonify({'success': False, 'message': 'Backup não fornecido'}), 400
        backup = data.get('backup') if isinstance(data, dict) else data
        if not isinstance(backup, list):
            return jsonify({'success': False, 'message': 'Backup deve ser uma lista de pacientes'}), 400
        erros = []
        for idx, registro in enumerate(backup):
            if not isinstance(registro, dict):
                erros.append(f'Item {idx + 1} não é um objeto válido')
                continue
            identificacao = registro.get('identificacao', {})
            avaliacao = registro.get('avaliacao', {})
            if not identificacao.get('nome_gestante'):
                erros.append(f'Item {idx + 1} não possui nome da gestante')
            if not isinstance(avaliacao, dict) or 'consultas_pre_natal' not in avaliacao:
                erros.append(f'Item {idx + 1} com avaliação incompleta')
        if erros:
            return jsonify({'success': False, 'message': 'Backup inválido', 'errors': erros}), 400
        return jsonify({'success': True, 'message': 'Backup válido', 'total': len(backup)})
    except Exception as e:
        return jsonify({'success': False, 'message': f'Erro ao validar backup: {str(e)}'}), 500

@app.route('/api/backup/limpar', methods=['DELETE'])
def limpar_banco_dados():
    """Remove todos os dados do banco de dados"""
    try:
        resultado = db.limpar_todos_dados()
        return jsonify(resultado)
    except Exception as e:
        return jsonify({'success': False, 'message': f'Erro ao limpar banco de dados: {str(e)}'}), 500

@app.route('/api/indicadores')
def indicadores():
    """Calcula os indicadores agregados dos pacientes usando o banco de dados"""
    try:
        unidade_saude = request.args.get('unidade_saude')
        stats = db.obter_estatisticas(unidade_saude=unidade_saude)
        return jsonify({
            'total': stats['total_pacientes'],
            'inicio_pre_natal_antes_12s': stats['inicio_pre_natal_antes_12s'],
            'consultas_pre_natal': stats['consultas_pre_natal'],
            'vacinas_completas': stats['vacinas_completas'],
            'plano_parto': stats['plano_parto'],
            'participou_grupos': stats['participou_grupos']
        })
    except Exception as e:
        return jsonify({
            'total': 0,
            'inicio_pre_natal_antes_12s': {'sim': 0, 'nao': 0},
            'consultas_pre_natal': {'ate_6': 0, 'mais_6': 0},
            'vacinas_completas': {'completa': 0, 'incompleta': 0, 'nao_avaliado': 0},
            'plano_parto': {'sim': 0, 'nao': 0},
            'participou_grupos': {'sim': 0, 'nao': 0}
        }), 500

@app.route('/api/unidades_saude')
def listar_unidades_saude():
    """Lista todas as unidades de saúde únicas"""
    try:
        unidades = db.obter_unidades_saude_unicas()
        return jsonify({
            'success': True,
            'unidades': unidades
        })
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/ranking_unidades')
def ranking_unidades():
    """Retorna ranking das unidades por indicador (percentual positivo), para mini-dashboard"""
    try:
        criterio = request.args.get('criterio', 'inicio_pre_natal_antes_12s')
        limite_raw = int(request.args.get('limite', 10))
        limite = 0 if limite_raw <= 0 else min(limite_raw, 50)
        unidades_param = (request.args.get('unidades') or '').strip()
        criterios_validos = {
            'inicio_pre_natal_antes_12s': ('sim', 'nao', 'Início pré-natal antes de 12 sem.'),
            'consultas_pre_natal': ('mais_6', 'ate_6', '≥ 6 consultas'),
            'vacinas_completas': ('completa', ('incompleta', 'nao_avaliado'), 'Vacinas completas'),
            'plano_parto': ('sim', 'nao', 'Plano de parto'),
            'participou_grupos': ('sim', 'nao', 'Participou de grupos')
        }
        criterio_hardcoded = criterio in criterios_validos
        if criterio_hardcoded:
            cfg = criterios_validos[criterio]
            pos_key = cfg[0]
            neg_keys = cfg[1] if isinstance(cfg[1], (list, tuple)) else (cfg[1],)
            label = cfg[2]
        else:
            # Validar coluna dinâmica antes de iterar
            coluna_check = db.obter_estatisticas_coluna(criterio)
            if not coluna_check.get('success'):
                return jsonify({'success': False, 'message': coluna_check.get('message', 'Filtro inválido'), 'ranking': []}), 400
            label = criterio

        unidades = db.obter_unidades_saude_unicas()
        if unidades_param:
            req_list = [x.strip() for x in unidades_param.split(',') if x.strip()]
            req_set = set((x or '').lower() for x in req_list)
            unidades = [u for u in unidades if (u or '').strip().lower() in req_set]
        ranking = []
        for u in unidades:
            if criterio_hardcoded:
                s = db.obter_estatisticas(unidade_saude=u)
                total = s['total_pacientes']
                if criterio == 'vacinas_completas':
                    d = s['vacinas_completas']
                    pos = d.get('completa', 0)
                    den = d.get('completa', 0) + d.get('incompleta', 0) + d.get('nao_avaliado', 0)
                    pct = (pos / den * 100) if den else 0
                elif criterio == 'consultas_pre_natal':
                    d = s['consultas_pre_natal']
                    pos = d.get('mais_6', 0)
                    den = d.get('mais_6', 0) + d.get('ate_6', 0)
                    pct = (pos / den * 100) if den else 0
                else:
                    d = s.get(criterio, {})
                    pos = d.get(pos_key, 0)
                    den = pos + sum(d.get(k, 0) for k in neg_keys)
                    pct = (pos / den * 100) if den else 0
            else:
                stats = db.obter_estatisticas_coluna(criterio, unidade_saude=u)
                pos = stats['data'].get('sim', 0)
                den = pos + stats['data'].get('nao', 0)
                total = den
                pct = (pos / den * 100) if den else 0
            ranking.append({
                'unidade': u,
                'total': total,
                'percentual': round(pct, 1),
                'criterio_label': label
            })
        ranking.sort(key=lambda x: (-x['percentual'], -x['total']))
        if unidades_param or limite <= 0:
            resultado_ranking = ranking
        else:
            resultado_ranking = ranking[:limite]
        return jsonify({
            'success': True,
            'criterio': criterio,
            'criterio_label': label,
            'ranking': resultado_ranking
        })
    except Exception as e:
        return jsonify({'success': False, 'message': str(e), 'ranking': []}), 500

@app.route('/api/indicadores_coluna/<nome_coluna>', methods=['GET'])
def obter_indicador_coluna(nome_coluna):
    """Retorna estatísticas de uma coluna genérica do BD"""
    try:
        unidade_saude = request.args.get('unidade_saude')
        resultado = db.obter_estatisticas_coluna(nome_coluna, unidade_saude)
        if resultado['success']:
            return jsonify({
                'success': True,
                'data': resultado['data']
            })
        else:
            return jsonify({
                'success': False,
                'message': resultado.get('message', 'Erro ao processar coluna'),
                'data': resultado['data']
            }), 400
    except Exception as e:
        return jsonify({
            'success': False,
            'message': str(e),
            'data': {'sim': 0, 'nao': 0}
        }), 500

@app.route('/api/campos_disponiveis', methods=['GET'])
def listar_campos_disponiveis():
    """Lista todos os campos disponíveis no BD (exceto nome e unidade_saude)"""
    try:
        campos_excluidos = {'id', 'nome_gestante', 'unidade_saude', 'data_salvamento', 'arquivo_origem'}
        
        # Obter nomes das colunas da tabela pacientes
        cursor = db.conn.cursor()
        cursor.execute("PRAGMA table_info(pacientes)")
        colunas = cursor.fetchall()
        
        # Se não houver colunas, retornar erro
        if not colunas:
            return jsonify({
                'success': False,
                'message': 'Nenhuma coluna encontrada na tabela pacientes',
                'campos': []
            }), 500
        
        # Mapear nomes técnicos para nomes amigáveis
        nomes_amigaveis = {
            'inicio_pre_natal_antes_12s': 'Início pré-natal antes de 12 semanas',
            'inicio_pre_natal_semanas': 'Semanas do início do pré-natal',
            'inicio_pre_natal_observacao': 'Observação do início do pré-natal',
            'consultas_pre_natal': 'Consultas de pré-natal',
            'vacinas_completas': 'Vacinas completas',
            'plano_parto': 'Plano de parto',
            'participou_grupos': 'Participou de grupos',
            'avaliacao_odontologica': 'Avaliação odontológica',
            'estratificacao': 'Estratificação',
            'estratificacao_problema': 'Problema de estratificação',
            'cartao_pre_natal_completo': 'Cartão pré-natal completo',
            'dum': 'DUM (Data da Última Menstruação)',
            'dpp': 'DPP (Data Provável do Parto)',
            'ganhou_kit': 'Ganhou kit',
            'kit_tipo': 'Tipo de kit',
            'proxima_avaliacao': 'Próxima avaliação',
            'proxima_avaliacao_hora': 'Hora da próxima avaliação',
            'ja_ganhou_crianca': 'Já ganhou criança',
            'data_ganhou_crianca': 'Data que ganhou criança',
            'quantidade_filhos': 'Quantidade de filhos',
            'generos_filhos': 'Gêneros dos filhos',
            'metodo_preventivo': 'Método preventivo',
            'metodo_preventivo_outros': 'Outro método preventivo'
        }
        
        campos_disponiveis = []
        # PRAGMA table_info retorna: cid, name, type, notnull, dflt_value, pk
        for coluna in colunas:
            # sqlite3.Row permite acesso como dicionário
            nome_coluna = coluna['name']
            tipo_coluna = coluna['type']
            
            if nome_coluna not in campos_excluidos:
                campo_info = {
                    'campo': nome_coluna,
                    'label': nomes_amigaveis.get(nome_coluna, nome_coluna.replace('_', ' ').title()),
                    'tipo': tipo_coluna
                }
                campos_disponiveis.append(campo_info)
        
        return jsonify({
            'success': True,
            'campos': campos_disponiveis
        })
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({
            'success': False,
            'message': f'Erro ao buscar campos: {str(e)}',
            'campos': []
        }), 500

@app.route('/api/indicadores/temporais/<filtro>', methods=['GET'])
def indicadores_temporais(filtro):
    """Retorna estatísticas temporais agrupadas por data para um indicador específico"""
    try:
        # Validar filtro
        filtros_validos = ['inicio_pre_natal_antes_12s', 'consultas_pre_natal',
                          'vacinas_completas', 'plano_parto', 'participou_grupos']
        if filtro not in filtros_validos:
            return jsonify({'error': 'Filtro inválido'}), 400

        stats_temporais = db.obter_estatisticas_temporais(filtro)
        return jsonify(stats_temporais)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

# ========== ROTAS DE GERENCIAMENTO DE TEMAS ==========

@app.route('/api/tema/obter', methods=['GET'])
def obter_tema():
    """Retorna o tema atual salvo no cookie do usuário"""
    try:
        # Como os cookies são gerenciados no frontend, retornamos apenas confirmação
        # O tema real é obtido via JavaScript no navegador
        return jsonify({
            'success': True,
            'message': 'Tema obtido via cookie do navegador'
        })
    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Erro ao obter tema: {str(e)}'
        }), 500

@app.route('/api/tema/salvar', methods=['POST'])
def salvar_tema():
    """Endpoint para informações sobre salvamento de tema (o salvamento real é feito via cookie)"""
    try:
        data = request.get_json()

        if not data:
            return jsonify({'success': False, 'message': 'Dados do tema não fornecidos'}), 400

        # Validação básica dos dados
        if 'nome' not in data or 'cores' not in data:
            return jsonify({'success': False, 'message': 'Estrutura do tema inválida'}), 400

        # Validar que cores é um objeto
        if not isinstance(data['cores'], dict):
            return jsonify({'success': False, 'message': 'Campo cores deve ser um objeto'}), 400

        # Validar formato das cores (hexadecimais)
        import re
        for key, color in data['cores'].items():
            if not re.match(r'^#[0-9A-Fa-f]{6}$', color):
                return jsonify({
                    'success': False,
                    'message': f'Cor inválida para {key}: {color}. Deve ser hexadecimal (#RRGGBB)'
                }), 400

        # Limitar tamanho do JSON (aproximadamente 4KB para cookie)
        import json
        theme_json = json.dumps(data)
        if len(theme_json) > 4000:
            return jsonify({'success': False, 'message': 'Tema muito grande para salvar'}), 400

        return jsonify({
            'success': True,
            'message': 'Tema validado com sucesso. Salvamento feito via cookie no navegador.',
            'tema': {
                'nome': data['nome'],
                'cores_count': len(data['cores'])
            }
        })

    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Erro ao processar tema: {str(e)}'
        }), 500

@app.route('/api/tema/padrao', methods=['GET'])
def obter_tema_padrao():
    """Retorna os valores padrão de cores do variables.css"""
    try:
        tema_padrao = {
            "nome": "Padrão Hospitalar",
            "cores": {
                "--bg": "#f2f5f4",
                "--card": "#ffffff",
                "--primary": "#2f7d6d",
                "--primary-hover": "#266758",
                "--secondary": "#4a90a4",
                "--secondary-hover": "#3d7a8a",
                "--accent": "#e6f2ef",
                "--text": "#263238",
                "--text-muted": "#607d8b",
                "--success": "#2e7d32",
                "--success-light": "#c8e6c9",
                "--warning": "#f57c00",
                "--warning-light": "#ffe0b2",
                "--danger": "#c62828",
                "--danger-light": "#ffcdd2",
                "--info": "#0277bd",
                "--info-light": "#b3e5fc",
                "--border": "#d0d7d5",
                "--border-light": "#e8edec"
            }
        }

        return jsonify({
            'success': True,
            'tema': tema_padrao
        })

    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Erro ao obter tema padrão: {str(e)}'
        }), 500

@app.route('/api/tema/salvar_css', methods=['POST'])
def salvar_css():
    """Salva um arquivo CSS customizado na pasta static/css"""
    try:
        data = request.get_json()
        
        if not data or 'css' not in data:
            return jsonify({'success': False, 'message': 'CSS não fornecido'}), 400
        
        css_content = data['css']
        
        # Validar que é CSS válido
        if not css_content.strip():
            return jsonify({'success': False, 'message': 'CSS vazio'}), 400
        
        # Gerar nome aleatório para o arquivo
        import uuid
        import hashlib
        css_hash = hashlib.md5(css_content.encode('utf-8')).hexdigest()[:8]
        filename = f'custom_{css_hash}.css'
        
        # Caminho completo do arquivo
        css_dir = os.path.join(os.path.dirname(__file__), 'static', 'css')
        os.makedirs(css_dir, exist_ok=True)
        file_path = os.path.join(css_dir, filename)
        
        # Salvar arquivo
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(css_content)
        
        return jsonify({
            'success': True,
            'filename': filename,
            'message': f'CSS salvo como {filename}'
        })
        
    except Exception as e:
        import traceback
        return jsonify({
            'success': False,
            'message': f'Erro ao salvar CSS: {str(e)}',
            'traceback': traceback.format_exc()
        }), 500

@app.route('/api/exportar/<formato>', methods=['GET'])
def exportar_pacientes(formato):
    """Exporta os dados dos pacientes no formato especificado"""
    try:
        # Buscar todos os pacientes inicialmente
        pacientes = db.obter_todos_pacientes()
        
        # Aplicar filtros de UPAs (múltiplas)
        upas = request.args.getlist('upas')
        if upas:
            pacientes = [p for p in pacientes if p.get('identificacao', {}).get('unidade_saude') in upas]
        
        # Aplicar filtros de Campos (múltiplos) - funciona em conjunto com UPAs (AND)
        campos = request.args.getlist('campos')
        if campos:
            pacientes = [p for p in pacientes if all(
                # Verificar se o campo tem valor não-nulo e não-vazio
                # Buscar em avaliacao primeiro (maioria dos campos)
                (lambda campo=campo: (
                    # Verificar em avaliacao
                    (p.get('avaliacao', {}).get(campo) is not None and 
                     p.get('avaliacao', {}).get(campo) != '') or
                    # Verificar em identificacao
                    (p.get('identificacao', {}).get(campo) is not None and 
                     p.get('identificacao', {}).get(campo) != '') or
                    # Verificar na raiz (casos especiais como id, data_salvamento)
                    (p.get(campo) is not None and p.get(campo) != '')
                ))() for campo in campos
            )]
        
        # Manter compatibilidade com filtro único antigo (deprecated)
        unidade_saude = request.args.get('unidade_saude')
        if unidade_saude and not upas and not campos:
            filtro = {'unidade_saude': unidade_saude}
            pacientes = db.buscar_pacientes(filtro)
        
        filtros = {
            'inicio_pre_natal': request.args.get('inicio_pre_natal'),
            'plano_parto': request.args.get('plano_parto'),
            'participou_grupos': request.args.get('participou_grupos'),
            'vacinas': request.args.get('vacinas')
        }
        pacientes_filtrados = aplicar_filtros_exportacao(pacientes, filtros)
        colunas_personalizadas = request.args.getlist('colunas')
        colunas_personalizadas = colunas_personalizadas if colunas_personalizadas else None

        if formato == 'excel':
            return exportar_excel(pacientes_filtrados, colunas_personalizadas)
        elif formato == 'word':
            return exportar_word(pacientes_filtrados, colunas_personalizadas)
        elif formato == 'txt':
            return exportar_txt(pacientes_filtrados, colunas_personalizadas)
        else:
            return jsonify({'success': False, 'message': 'Formato não suportado'}), 400
            
    except Exception as e:
        return jsonify({'success': False, 'message': f'Erro ao exportar: {str(e)}'}), 500

def exportar_excel(pacientes, colunas_personalizadas=None):
    """Exporta dados para Excel (.xlsx)"""
    try:
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, PatternFill, Alignment
        except ImportError:
            return jsonify({
                'success': False,
                'message': 'Biblioteca openpyxl não encontrada. Instale com: pip install openpyxl'
            }), 500

        wb = Workbook()
        ws = wb.active
        ws.title = "Pacientes"

        colunas = obter_colunas_config(colunas_personalizadas)
        headers = [coluna['label'] for coluna in colunas]

        header_fill = PatternFill(start_color="667eea", end_color="764ba2", fill_type="solid")
        header_font = Font(bold=True, color="FFFFFF", size=12)
        header_alignment = Alignment(horizontal="center", vertical="center")

        for col_idx, header in enumerate(headers, start=1):
            cell = ws.cell(row=1, column=col_idx, value=header)
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = header_alignment

        for row_idx, paciente in enumerate(pacientes, start=2):
            for col_idx, coluna in enumerate(colunas, start=1):
                ws.cell(row=row_idx, column=col_idx, value=obter_valor_coluna(paciente, coluna))

        for col_idx in range(1, len(headers) + 1):
            ws.column_dimensions[ws.cell(row=1, column=col_idx).column_letter].width = 22

        output = io.BytesIO()
        wb.save(output)
        output.seek(0)

        filename = f"pacientes_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
        response = make_response(send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=filename
        ))
        return response

    except Exception as e:
        return jsonify({'success': False, 'message': f'Erro ao exportar Excel: {str(e)}'}), 500

def exportar_word(pacientes, colunas_personalizadas=None):
    """Exporta dados para Word (.docx)"""
    try:
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            return jsonify({
                'success': False,
                'message': 'Biblioteca python-docx não encontrada. Instale com: pip install python-docx'
            }), 500

        doc = Document()
        colunas = obter_colunas_config(colunas_personalizadas)

        title = doc.add_heading('Relatório de Pacientes', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f'Total de pacientes: {len(pacientes)}')
        doc.add_paragraph(f'Data de exportação: {datetime.now().strftime("%d/%m/%Y %H:%M:%S")}')
        doc.add_paragraph('')

        for paciente in pacientes:
            ident = paciente.get('identificacao', {})
            nome = ident.get('nome_gestante', 'Nome não informado')
            doc.add_heading(nome, level=2)

            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Campo'
            header_cells[1].text = 'Valor'

            for coluna in colunas:
                row_cells = table.add_row().cells
                row_cells[0].text = coluna['label']
                row_cells[1].text = str(obter_valor_coluna(paciente, coluna))

            doc.add_paragraph('')

        output = io.BytesIO()
        doc.save(output)
        output.seek(0)

        filename = f"pacientes_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        response = make_response(send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=filename
        ))
        return response

    except Exception as e:
        return jsonify({'success': False, 'message': f'Erro ao exportar Word: {str(e)}'}), 500

def exportar_txt(pacientes, colunas_personalizadas=None):
    """Exporta dados para arquivo de texto (.txt)"""
    try:
        output = io.StringIO()
        colunas = obter_colunas_config(colunas_personalizadas)

        output.write("=" * 80 + "\n")
        output.write("RELATÓRIO DE PACIENTES\n")
        output.write("=" * 80 + "\n")
        output.write(f"Total de pacientes: {len(pacientes)}\n")
        output.write(f"Data de exportação: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}\n")
        output.write("=" * 80 + "\n\n")

        for idx, paciente in enumerate(pacientes, start=1):
            output.write(f"\n{'=' * 80}\n")
            output.write(f"PACIENTE {idx}\n")
            output.write(f"{'=' * 80}\n\n")

            for coluna in colunas:
                output.write(f"  {coluna['label']}: {obter_valor_coluna(paciente, coluna)}\n")
            output.write("\n")

        content = output.getvalue()
        output.close()
        output_bytes = io.BytesIO(content.encode('utf-8'))

        filename = f"pacientes_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
        response = make_response(send_file(
            output_bytes,
            mimetype='text/plain; charset=utf-8',
            as_attachment=True,
            download_name=filename
        ))
        return response

    except Exception as e:
        return jsonify({'success': False, 'message': f'Erro ao exportar TXT: {str(e)}'}), 500

def aplicar_filtros_exportacao(pacientes, filtros):
    """Aplica os filtros selecionados antes da exportação"""
    if not any(filtros.values()):
        return pacientes

    resultados = []
    for paciente in pacientes:
        avaliacao = paciente.get('avaliacao', {})

        if filtros.get('inicio_pre_natal') and not avaliacao.get('inicio_pre_natal_antes_12s'):
            continue
        if filtros.get('plano_parto') and not avaliacao.get('plano_parto'):
            continue
        if filtros.get('participou_grupos') and not avaliacao.get('participou_grupos'):
            continue

        if filtros.get('vacinas'):
            status_vacina = (avaliacao.get('vacinas_completas', '') or '').lower()
            if not verificar_status_vacinas(status_vacina, filtros['vacinas']):
                continue

        resultados.append(paciente)

    return resultados

def verificar_status_vacinas(status, filtro):
    if filtro == 'completa':
        return 'complet' in status and 'incomplet' not in status
    if filtro == 'incompleta':
        return 'incomplet' in status
    if filtro == 'nao_avaliado':
        return status == '' or 'nao' in status or 'não' in status
    return True

def formatar_boolean(valor):
    """Formata valores booleanos para exibição"""
    if valor is True:
        return 'Sim'
    elif valor is False:
        return 'Não'
    return 'Não informado'

def formatar_vacinas(valor):
    if valor:
        return valor
    return 'Não avaliado'

def formatar_valor_opcional(valor):
    """Formata valores opcionais (None, vazio, etc.)"""
    if valor is None or valor == '':
        return 'Não informado'
    return str(valor)

COLUNAS_CONFIG = [
    {'key': 'identificacao.nome_gestante', 'label': 'Nome da Gestante'},
    {'key': 'identificacao.unidade_saude', 'label': 'Unidade de Saúde'},
    {'key': 'data_salvamento', 'label': 'Data de Cadastro'},
    {
        'key': 'avaliacao.inicio_pre_natal_antes_12s',
        'label': 'Início pré-natal antes de 12 semanas',
        'formatter': formatar_boolean
    },
    {'key': 'avaliacao.inicio_pre_natal_semanas', 'label': 'Semanas de início do pré-natal'},
    {'key': 'avaliacao.inicio_pre_natal_observacao', 'label': 'Observação do início pré-natal'},
    {'key': 'avaliacao.consultas_pre_natal', 'label': 'Consultas de pré-natal'},
    {
        'key': 'avaliacao.vacinas_completas',
        'label': 'Vacinas completas',
        'formatter': formatar_vacinas
    },
    {
        'key': 'avaliacao.plano_parto',
        'label': 'Plano de parto',
        'formatter': formatar_boolean
    },
    {
        'key': 'avaliacao.participou_grupos',
        'label': 'Participou de grupos',
        'formatter': formatar_boolean
    },
    {
        'key': 'avaliacao.avaliacao_odontologica',
        'label': 'Avaliação odontológica',
        'formatter': formatar_boolean
    },
    {
        'key': 'avaliacao.estratificacao',
        'label': 'Estratificação',
        'formatter': formatar_boolean
    },
    {'key': 'avaliacao.estratificacao_problema', 'label': 'Problema de estratificação'},
    {
        'key': 'avaliacao.cartao_pre_natal_completo',
        'label': 'Cartão pré-natal completo',
        'formatter': formatar_boolean
    },
    {'key': 'avaliacao.dum', 'label': 'DUM (Data da Última Menstruação)'},
    {'key': 'avaliacao.dpp', 'label': 'DPP (Data Provável do Parto)'},
    {
        'key': 'avaliacao.ganhou_kit',
        'label': 'Ganhou kit',
        'formatter': formatar_boolean
    },
    {'key': 'avaliacao.kit_tipo', 'label': 'Tipo de kit'},
    {'key': 'avaliacao.proxima_avaliacao', 'label': 'Próxima avaliação (data)'},
    {'key': 'avaliacao.proxima_avaliacao_hora', 'label': 'Próxima avaliação (hora)'},
    {
        'key': 'avaliacao.ja_ganhou_crianca',
        'label': 'Já ganhou criança',
        'formatter': formatar_boolean
    },
    {'key': 'avaliacao.data_ganhou_crianca', 'label': 'Data que ganhou criança'},
    {'key': 'avaliacao.quantidade_filhos', 'label': 'Quantidade de filhos'},
    {'key': 'avaliacao.generos_filhos', 'label': 'Gêneros dos filhos'},
    {'key': 'avaliacao.metodo_preventivo', 'label': 'Método preventivo'},
    {'key': 'avaliacao.metodo_preventivo_outros', 'label': 'Método preventivo (outros)'}
]

def obter_colunas_config(colunas_personalizadas):
    if colunas_personalizadas:
        selecionadas = [
            coluna for coluna in COLUNAS_CONFIG if coluna['key'] in colunas_personalizadas
        ]
        if selecionadas:
            return selecionadas
    return COLUNAS_CONFIG

def obter_valor_coluna(paciente, coluna):
    valor = paciente
    for parte in coluna['key'].split('.'):
        if isinstance(valor, dict):
            valor = valor.get(parte, '')
        else:
            valor = ''
    formatter = coluna.get('formatter')
    if callable(formatter):
        return formatter(valor)
    if valor is None:
        return ''
    return valor

# Variável global para armazenar os servidores Flask
_flask_server = None
_flask_servers = {}  # Dicionário para armazenar múltiplos servidores por porta

def _get_db_for_port(port):
    """Retorna a instância do banco de dados para uma porta específica"""
    global db, _db_instances
    
    if port not in _db_instances:
        # Configurar banco de dados baseado na porta
        if port == 5001:
            # Porta 5001: usar banco de dados copiado
            base_dir = os.path.dirname(os.path.abspath(__file__))
            db_path_porta5001 = os.path.join(base_dir, 'data', 'pacientes_porta5001.db')
            _db_instances[port] = Database(db_path=db_path_porta5001)
            print(f"Banco de dados configurado para porta {port}: {db_path_porta5001}")
        else:
            # Outras portas: usar banco de dados padrão
            _db_instances[port] = Database()
    
    return _db_instances[port]

def run_flask(debug=False, use_reloader=False, silent=False, port=None):
    """Inicia o servidor Flask"""
    global _flask_server, db, _thread_local
    
    # Obter configurações do .env ou usar padrões
    host = os.getenv('FLASK_HOST', '127.0.0.1')
    if port is None:
        port = int(os.getenv('PORT', 5000))
    
    # Configurar banco de dados para esta porta e thread
    db_instance = _get_db_for_port(port)
    _thread_local.db_instance = db_instance
    # Se debug não foi passado explicitamente, verificar .env
    if debug is False:
        debug = os.getenv('FLASK_DEBUG', 'false').lower() == 'true'
    
    # Verificar se está em modo executável e debug
    is_executable = getattr(sys, 'frozen', False)
    
    if silent or (is_executable and not debug):
        import logging
        log = logging.getLogger('werkzeug')
        log.setLevel(logging.ERROR)
    
    threads = int(os.getenv('WAITRESS_THREADS', '8'))
    
    # Descoberta: Zeroconf (padrão) ou scan /24 (DISCOVERY=scan)
    discovery = (os.getenv('DISCOVERY') or 'zeroconf').strip().lower()
    try:
        if discovery == 'scan':
            from inicio.rede import run_scan_loop
            run_scan_loop(port=port)
        else:
            from inicio.rede.zeroconf_discovery import start_zeroconf
            start_zeroconf(port)
    except Exception:
        pass

    if use_reloader:
        # Apenas em desenvolvimento: Flask com reloader (ignora Waitress)
        app.run(host=host, port=port, debug=debug, use_reloader=use_reloader)
    else:
        # Produção: Waitress — funciona em .py, .exe e em qualquer PC.
        # Ignora FLASK_RUN_HOST, flask run e defaults do Flask.
        from waitress import serve
        try:
            if not is_executable or debug:
                print(f"Servidor iniciado em http://{host}:{port} (Waitress, {threads} threads)")
            serve(app, host=host, port=port, threads=threads)
        except KeyboardInterrupt:
            if not is_executable or debug:
                print(f"\nServidor na porta {port} encerrado")
        except Exception as e:
            if is_executable and not debug:
                try:
                    import tkinter as tk
                    from tkinter import messagebox
                    root = tk.Tk()
                    root.withdraw()
                    messagebox.showerror(
                        "Erro ao Iniciar Servidor",
                        f"Não foi possível iniciar o servidor na porta {port}.\n\n"
                        f"Erro: {str(e)}\n\n"
                        f"Verifique se:\n"
                        f"- A porta não está em uso\n"
                        f"- Você tem permissões suficientes\n"
                        f"- O firewall não está bloqueando"
                    )
                    root.destroy()
                except Exception:
                    pass
            else:
                print(f"Erro ao iniciar servidor: {e}")
                import traceback
                traceback.print_exc()
            raise

def cleanup_flask():
    """Limpeza ao encerrar aplicação"""
    global _flask_server, _flask_servers
    # Encerrar todos os servidores
    for port, server in list(_flask_servers.items()):
        try:
            print(f"Encerrando servidor Flask na porta {port}...")
            server.shutdown()
            print(f"Servidor Flask na porta {port} encerrado")
        except:
            pass
    _flask_servers.clear()
    if _flask_server:
        try:
            print("Encerrando servidor Flask...")
            _flask_server.shutdown()
            _flask_server = None
            print("Servidor Flask encerrado")
        except:
            pass
    try:
        from inicio.rede.zeroconf_discovery import stop_zeroconf
        stop_zeroconf()
    except Exception:
        pass

# Registrar cleanup
atexit.register(cleanup_flask)
