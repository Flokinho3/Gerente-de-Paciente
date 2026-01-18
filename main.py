import threading
import webbrowser
import json
import os
import io
from flask import Flask, jsonify, render_template, request, send_file, make_response
from datetime import datetime
import tkinter as tk
from tkinter import messagebox
from database import db

app = Flask(__name__)

@app.route('/')
def home():
    return render_template('Home.html')

@app.route('/novo_paciente')
def novo_paciente():
    return render_template('novo_paciente.html')

@app.route('/pacientes')
def pacientes():
    return render_template('pacientes.html')

@app.route('/exportar')
def exportar():
    return render_template('exportar.html')

@app.route('/bd')
def bd():
    return render_template('bd.html')

@app.route('/api/salvar_paciente', methods=['POST'])
def salvar_paciente():
    """Salva os dados do paciente usando o banco de dados"""
    try:
        data = request.get_json()
        
        # Validar dados obrigatórios
        if not data or 'identificacao' not in data or 'avaliacao' not in data:
            return jsonify({'success': False, 'message': 'Dados inválidos'}), 400
        
        nome = data['identificacao'].get('nome_gestante', '').strip()
        if not nome:
            return jsonify({'success': False, 'message': 'Nome da gestante é obrigatório'}), 400
        
        # Salvar usando o banco de dados
        resultado = db.adicionar_paciente(data)
        
        return jsonify(resultado)
    except Exception as e:
        return jsonify({'success': False, 'message': f'Erro ao salvar: {str(e)}'}), 500

@app.route('/api/pacientes', methods=['GET'])
def listar_pacientes():
    """Lista todos os pacientes ou filtra por parâmetros"""
    try:
        filtro = {}
        
        # Filtros opcionais via query parameters
        nome = request.args.get('nome')
        unidade = request.args.get('unidade_saude')
        
        if nome:
            filtro['nome'] = nome
        if unidade:
            filtro['unidade_saude'] = unidade
        
        pacientes = db.buscar_pacientes(filtro if filtro else None)
        
        return jsonify({
            'success': True,
            'total': len(pacientes),
            'pacientes': pacientes
        })
    except Exception as e:
        return jsonify({'success': False, 'message': f'Erro ao buscar: {str(e)}'}), 500

@app.route('/api/atualizar_paciente/<paciente_id>', methods=['PUT'])
def atualizar_paciente(paciente_id):
    """Atualiza os dados de um paciente existente"""
    try:
        data = request.get_json()
        
        # Validar dados obrigatórios
        if not data or 'identificacao' not in data or 'avaliacao' not in data:
            return jsonify({'success': False, 'message': 'Dados inválidos'}), 400
        
        # Atualizar usando o banco de dados
        resultado = db.atualizar_paciente(paciente_id, data)
        
        if resultado['success']:
            return jsonify(resultado)
        else:
            return jsonify(resultado), 404
        
    except Exception as e:
        return jsonify({'success': False, 'message': f'Erro ao atualizar: {str(e)}'}), 500

@app.route('/api/deletar_paciente/<paciente_id>', methods=['DELETE'])
def deletar_paciente(paciente_id):
    """Deleta um paciente do banco de dados"""
    try:
        resultado = db.deletar_paciente(paciente_id)
        
        if resultado['success']:
            return jsonify(resultado)
        else:
            return jsonify(resultado), 404
            
    except Exception as e:
        return jsonify({'success': False, 'message': f'Erro ao deletar: {str(e)}'}), 500

@app.route('/api/backup/criar', methods=['GET'])
def criar_backup():
    """Cria um backup completo do banco de dados"""
    try:
        resultado = db.criar_backup()
        return jsonify(resultado)
    except Exception as e:
        return jsonify({'success': False, 'message': f'Erro ao criar backup: {str(e)}'}), 500

@app.route('/api/backup/download', methods=['GET'])
def download_backup():
    """Permite baixar o último backup como arquivo JSON"""
    try:
        resultado = db.criar_backup()
        buffer = io.BytesIO(json.dumps(resultado['backup'], ensure_ascii=False, indent=2).encode('utf-8'))
        buffer.seek(0)
        filename = f"backup_pacientes_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
        response = make_response(send_file(
            buffer,
            mimetype='application/json',
            as_attachment=True,
            download_name=filename
        ))
        return response
    except Exception as e:
        return jsonify({'success': False, 'message': f'Erro ao gerar arquivo de backup: {str(e)}'}), 500

@app.route('/api/backup/restaurar', methods=['POST'])
def restaurar_backup():
    """Restaura o banco de dados a partir de um backup"""
    try:
        data = request.get_json()
        if not data:
            return jsonify({'success': False, 'message': 'Dados do backup não fornecidos'}), 400

        backup = data.get('backup') if isinstance(data, dict) else data
        if not isinstance(backup, list):
            return jsonify({'success': False, 'message': 'Backup deve ser uma lista de registros'}), 400

        resultado = db.restaurar_backup(backup)
        
        if resultado['success']:
            return jsonify(resultado)
        else:
            return jsonify(resultado), 400
            
    except Exception as e:
        return jsonify({'success': False, 'message': f'Erro ao restaurar backup: {str(e)}'}), 500

@app.route('/api/backup/validar', methods=['POST'])
def validar_backup():
    """Valida a estrutura do backup antes de restaurar"""
    try:
        data = request.get_json()
        if not data:
            return jsonify({'success': False, 'message': 'Backup não fornecido'}), 400
        backup = data.get('backup') if isinstance(data, dict) else data
        if not isinstance(backup, list):
            return jsonify({'success': False, 'message': 'Backup deve ser uma lista de pacientes'}), 400
        erros = []
        for idx, registro in enumerate(backup):
            if not isinstance(registro, dict):
                erros.append(f'Item {idx + 1} não é um objeto válido')
                continue
            identificacao = registro.get('identificacao', {})
            avaliacao = registro.get('avaliacao', {})
            if not identificacao.get('nome_gestante'):
                erros.append(f'Item {idx + 1} não possui nome da gestante')
            if not isinstance(avaliacao, dict) or 'consultas_pre_natal' not in avaliacao:
                erros.append(f'Item {idx + 1} com avaliação incompleta')
        if erros:
            return jsonify({'success': False, 'message': 'Backup inválido', 'errors': erros}), 400
        return jsonify({'success': True, 'message': 'Backup válido', 'total': len(backup)})
    except Exception as e:
        return jsonify({'success': False, 'message': f'Erro ao validar backup: {str(e)}'}), 500

@app.route('/api/backup/limpar', methods=['DELETE'])
def limpar_banco_dados():
    """Remove todos os dados do banco de dados"""
    try:
        resultado = db.limpar_todos_dados()
        return jsonify(resultado)
    except Exception as e:
        return jsonify({'success': False, 'message': f'Erro ao limpar banco de dados: {str(e)}'}), 500

@app.route('/api/indicadores')
def indicadores():
    """Calcula os indicadores agregados dos pacientes usando o banco de dados"""
    try:
        stats = db.obter_estatisticas()
        return jsonify({
            'total': stats['total_pacientes'],
            'inicio_pre_natal_antes_12s': stats['inicio_pre_natal_antes_12s'],
            'consultas_pre_natal': stats['consultas_pre_natal'],
            'vacinas_completas': stats['vacinas_completas'],
            'plano_parto': stats['plano_parto'],
            'participou_grupos': stats['participou_grupos']
        })
    except Exception as e:
        return jsonify({
            'total': 0,
            'inicio_pre_natal_antes_12s': {'sim': 0, 'nao': 0},
            'consultas_pre_natal': {'ate_6': 0, 'mais_6': 0},
            'vacinas_completas': {'completa': 0, 'incompleta': 0, 'nao_avaliado': 0},
            'plano_parto': {'sim': 0, 'nao': 0},
            'participou_grupos': {'sim': 0, 'nao': 0}
        }), 500

@app.route('/api/indicadores/temporais/<filtro>', methods=['GET'])
def indicadores_temporais(filtro):
    """Retorna estatísticas temporais agrupadas por data para um indicador específico"""
    try:
        # Validar filtro
        filtros_validos = ['inicio_pre_natal_antes_12s', 'consultas_pre_natal', 
                          'vacinas_completas', 'plano_parto', 'participou_grupos']
        if filtro not in filtros_validos:
            return jsonify({'error': 'Filtro inválido'}), 400
        
        stats_temporais = db.obter_estatisticas_temporais(filtro)
        return jsonify(stats_temporais)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/exportar/<formato>', methods=['GET'])
def exportar_pacientes(formato):
    """Exporta os dados dos pacientes no formato especificado"""
    try:
        pacientes = db.obter_todos_pacientes()
        filtros = {
            'inicio_pre_natal': request.args.get('inicio_pre_natal'),
            'plano_parto': request.args.get('plano_parto'),
            'participou_grupos': request.args.get('participou_grupos'),
            'vacinas': request.args.get('vacinas')
        }
        pacientes_filtrados = aplicar_filtros_exportacao(pacientes, filtros)
        colunas_personalizadas = request.args.getlist('colunas')
        colunas_personalizadas = colunas_personalizadas if colunas_personalizadas else None

        if formato == 'excel':
            return exportar_excel(pacientes_filtrados, colunas_personalizadas)
        elif formato == 'word':
            return exportar_word(pacientes_filtrados, colunas_personalizadas)
        elif formato == 'txt':
            return exportar_txt(pacientes_filtrados, colunas_personalizadas)
        else:
            return jsonify({'success': False, 'message': 'Formato não suportado'}), 400
            
    except Exception as e:
        return jsonify({'success': False, 'message': f'Erro ao exportar: {str(e)}'}), 500

def exportar_excel(pacientes, colunas_personalizadas=None):
    """Exporta dados para Excel (.xlsx)"""
    try:
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, PatternFill, Alignment
        except ImportError:
            return jsonify({
                'success': False,
                'message': 'Biblioteca openpyxl não encontrada. Instale com: pip install openpyxl'
            }), 500

        wb = Workbook()
        ws = wb.active
        ws.title = "Pacientes"

        colunas = obter_colunas_config(colunas_personalizadas)
        headers = [coluna['label'] for coluna in colunas]

        header_fill = PatternFill(start_color="667eea", end_color="764ba2", fill_type="solid")
        header_font = Font(bold=True, color="FFFFFF", size=12)
        header_alignment = Alignment(horizontal="center", vertical="center")

        for col_idx, header in enumerate(headers, start=1):
            cell = ws.cell(row=1, column=col_idx, value=header)
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = header_alignment

        for row_idx, paciente in enumerate(pacientes, start=2):
            for col_idx, coluna in enumerate(colunas, start=1):
                ws.cell(row=row_idx, column=col_idx, value=obter_valor_coluna(paciente, coluna))

        for col_idx in range(1, len(headers) + 1):
            ws.column_dimensions[ws.cell(row=1, column=col_idx).column_letter].width = 22

        output = io.BytesIO()
        wb.save(output)
        output.seek(0)

        filename = f"pacientes_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
        response = make_response(send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=filename
        ))
        return response

    except Exception as e:
        return jsonify({'success': False, 'message': f'Erro ao exportar Excel: {str(e)}'}), 500

def exportar_word(pacientes, colunas_personalizadas=None):
    """Exporta dados para Word (.docx)"""
    try:
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            return jsonify({
                'success': False,
                'message': 'Biblioteca python-docx não encontrada. Instale com: pip install python-docx'
            }), 500

        doc = Document()
        colunas = obter_colunas_config(colunas_personalizadas)

        title = doc.add_heading('Relatório de Pacientes', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f'Total de pacientes: {len(pacientes)}')
        doc.add_paragraph(f'Data de exportação: {datetime.now().strftime("%d/%m/%Y %H:%M:%S")}')
        doc.add_paragraph('')

        for paciente in pacientes:
            ident = paciente.get('identificacao', {})
            nome = ident.get('nome_gestante', 'Nome não informado')
            doc.add_heading(nome, level=2)

            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Campo'
            header_cells[1].text = 'Valor'

            for coluna in colunas:
                row_cells = table.add_row().cells
                row_cells[0].text = coluna['label']
                row_cells[1].text = str(obter_valor_coluna(paciente, coluna))

            doc.add_paragraph('')

        output = io.BytesIO()
        doc.save(output)
        output.seek(0)

        filename = f"pacientes_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        response = make_response(send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=filename
        ))
        return response

    except Exception as e:
        return jsonify({'success': False, 'message': f'Erro ao exportar Word: {str(e)}'}), 500

def exportar_txt(pacientes, colunas_personalizadas=None):
    """Exporta dados para arquivo de texto (.txt)"""
    try:
        output = io.StringIO()
        colunas = obter_colunas_config(colunas_personalizadas)

        output.write("=" * 80 + "\n")
        output.write("RELATÓRIO DE PACIENTES\n")
        output.write("=" * 80 + "\n")
        output.write(f"Total de pacientes: {len(pacientes)}\n")
        output.write(f"Data de exportação: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}\n")
        output.write("=" * 80 + "\n\n")

        for idx, paciente in enumerate(pacientes, start=1):
            output.write(f"\n{'=' * 80}\n")
            output.write(f"PACIENTE {idx}\n")
            output.write(f"{'=' * 80}\n\n")

            for coluna in colunas:
                output.write(f"  {coluna['label']}: {obter_valor_coluna(paciente, coluna)}\n")
            output.write("\n")

        content = output.getvalue()
        output.close()
        output_bytes = io.BytesIO(content.encode('utf-8'))

        filename = f"pacientes_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
        response = make_response(send_file(
            output_bytes,
            mimetype='text/plain; charset=utf-8',
            as_attachment=True,
            download_name=filename
        ))
        return response

    except Exception as e:
        return jsonify({'success': False, 'message': f'Erro ao exportar TXT: {str(e)}'}), 500

def aplicar_filtros_exportacao(pacientes, filtros):
    """Aplica os filtros selecionados antes da exportação"""
    if not any(filtros.values()):
        return pacientes

    resultados = []
    for paciente in pacientes:
        avaliacao = paciente.get('avaliacao', {})

        if filtros.get('inicio_pre_natal') and not avaliacao.get('inicio_pre_natal_antes_12s'):
            continue
        if filtros.get('plano_parto') and not avaliacao.get('plano_parto'):
            continue
        if filtros.get('participou_grupos') and not avaliacao.get('participou_grupos'):
            continue

        if filtros.get('vacinas'):
            status_vacina = (avaliacao.get('vacinas_completas', '') or '').lower()
            if not verificar_status_vacinas(status_vacina, filtros['vacinas']):
                continue

        resultados.append(paciente)

    return resultados


def verificar_status_vacinas(status, filtro):
    if filtro == 'completa':
        return 'complet' in status and 'incomplet' not in status
    if filtro == 'incompleta':
        return 'incomplet' in status
    if filtro == 'nao_avaliado':
        return status == '' or 'nao' in status or 'não' in status
    return True


def formatar_boolean(valor):
    """Formata valores booleanos para exibição"""
    if valor is True:
        return 'Sim'
    elif valor is False:
        return 'Não'
    return 'Não informado'


def formatar_vacinas(valor):
    if valor:
        return valor
    return 'Não avaliado'


COLUNAS_CONFIG = [
    {'key': 'identificacao.nome_gestante', 'label': 'Nome da Gestante'},
    {'key': 'identificacao.unidade_saude', 'label': 'Unidade de Saúde'},
    {'key': 'data_salvamento', 'label': 'Data de Cadastro'},
    {
        'key': 'avaliacao.inicio_pre_natal_antes_12s',
        'label': 'Início pré-natal antes de 12 semanas',
        'formatter': formatar_boolean
    },
    {'key': 'avaliacao.consultas_pre_natal', 'label': 'Consultas de pré-natal'},
    {
        'key': 'avaliacao.vacinas_completas',
        'label': 'Vacinas completas',
        'formatter': formatar_vacinas
    },
    {
        'key': 'avaliacao.plano_parto',
        'label': 'Plano de parto',
        'formatter': formatar_boolean
    },
    {
        'key': 'avaliacao.participou_grupos',
        'label': 'Participou de grupos',
        'formatter': formatar_boolean
    },
    {
        'key': 'avaliacao.avaliacao_odontologica',
        'label': 'Avaliação odontológica',
        'formatter': formatar_boolean
    },
    {
        'key': 'avaliacao.estratificacao',
        'label': 'Estratificação',
        'formatter': formatar_boolean
    },
    {
        'key': 'avaliacao.cartao_pre_natal_completo',
        'label': 'Cartão pré-natal completo',
        'formatter': formatar_boolean
    }
]


def obter_colunas_config(colunas_personalizadas):
    if colunas_personalizadas:
        selecionadas = [
            coluna for coluna in COLUNAS_CONFIG if coluna['key'] in colunas_personalizadas
        ]
        if selecionadas:
            return selecionadas
    return COLUNAS_CONFIG


def obter_valor_coluna(paciente, coluna):
    valor = paciente
    for parte in coluna['key'].split('.'):
        if isinstance(valor, dict):
            valor = valor.get(parte, '')
        else:
            valor = ''
    formatter = coluna.get('formatter')
    if callable(formatter):
        return formatter(valor)
    if valor is None:
        return ''
    return valor

def run_flask():
    app.run(host='127.0.0.1', port=5000, debug=True)

def main():

    """
    para ativar o debug do flask, basta comentar a linha 20 e descomentar a linha 21
    root = tk.Tk()
    root.withdraw()  # esconde a janela principal

    messagebox.showinfo(
        "Sistema de Gestão de Pacientes",
        "Sistema iniciado com sucesso.\nClique em OK para abrir no navegador."
    )

    threading.Thread(target=run_flask, daemon=True).start()
    webbrowser.open('http://localhost:5000')
    """
    # inicia sem o tkinter
    run_flask()

if __name__ == '__main__':
    main()
