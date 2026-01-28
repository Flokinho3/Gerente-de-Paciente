"""
API: Exportação de pacientes (Excel, Word, TXT).
"""
import io
from datetime import datetime

from flask import Blueprint, jsonify, make_response, request, send_file

from .db import db
from .exportar_helpers import (
    aplicar_filtros_exportacao,
    obter_colunas_config,
    obter_valor_coluna,
)

bp = Blueprint("api_exportar", __name__, url_prefix="/api")


def _exportar_excel(pacientes, colunas_personalizadas=None):
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Alignment, Font, PatternFill
    except ImportError:
        return jsonify({
            "success": False,
            "message": "Biblioteca openpyxl não encontrada. Instale com: pip install openpyxl",
        }), 500
    colunas = obter_colunas_config(colunas_personalizadas)
    headers = [c["label"] for c in colunas]
    wb = Workbook()
    ws = wb.active
    ws.title = "Pacientes"
    fill = PatternFill(start_color="667eea", end_color="764ba2", fill_type="solid")
    font = Font(bold=True, color="FFFFFF", size=12)
    align = Alignment(horizontal="center", vertical="center")
    for i, h in enumerate(headers, 1):
        c = ws.cell(row=1, column=i, value=h)
        c.fill, c.font, c.alignment = fill, font, align
    for ri, paciente in enumerate(pacientes, 2):
        for ci, col in enumerate(colunas, 1):
            ws.cell(row=ri, column=ci, value=obter_valor_coluna(paciente, col))
    for i in range(1, len(headers) + 1):
        ws.column_dimensions[ws.cell(row=1, column=i).column_letter].width = 22
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    name = f"pacientes_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    return make_response(
        send_file(buf, mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", as_attachment=True, download_name=name)
    )


def _exportar_word(pacientes, colunas_personalizadas=None):
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return jsonify({
            "success": False,
            "message": "Biblioteca python-docx não encontrada. Instale com: pip install python-docx",
        }), 500
    colunas = obter_colunas_config(colunas_personalizadas)
    doc = Document()
    t = doc.add_heading("Relatório de Pacientes", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Total de pacientes: {len(pacientes)}")
    doc.add_paragraph(f"Data de exportação: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
    doc.add_paragraph("")
    for paciente in pacientes:
        nome = paciente.get("identificacao", {}).get("nome_gestante", "Nome não informado")
        doc.add_heading(nome, level=2)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        table.rows[0].cells[0].text, table.rows[0].cells[1].text = "Campo", "Valor"
        for col in colunas:
            row = table.add_row().cells
            row[0].text, row[1].text = col["label"], str(obter_valor_coluna(paciente, col))
        doc.add_paragraph("")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    name = f"pacientes_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    return make_response(
        send_file(
            buf,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name=name,
        )
    )


def _exportar_txt(pacientes, colunas_personalizadas=None):
    colunas = obter_colunas_config(colunas_personalizadas)
    out = io.StringIO()
    out.write("=" * 80 + "\nRELATÓRIO DE PACIENTES\n" + "=" * 80 + "\n")
    out.write(f"Total de pacientes: {len(pacientes)}\n")
    out.write(f"Data de exportação: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}\n" + "=" * 80 + "\n\n")
    for idx, p in enumerate(pacientes, 1):
        out.write("\n" + "=" * 80 + f"\nPACIENTE {idx}\n" + "=" * 80 + "\n\n")
        for col in colunas:
            out.write(f"  {col['label']}: {obter_valor_coluna(p, col)}\n")
        out.write("\n")
    raw = out.getvalue().encode("utf-8")
    out.close()
    buf = io.BytesIO(raw)
    name = f"pacientes_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
    return make_response(
        send_file(buf, mimetype="text/plain; charset=utf-8", as_attachment=True, download_name=name)
    )


@bp.route("/exportar/<formato>", methods=["GET"])
def exportar_pacientes(formato):
    try:
        pacientes = db.obter_todos_pacientes()
        upas = request.args.getlist("upas")
        if upas:
            pacientes = [p for p in pacientes if p.get("identificacao", {}).get("unidade_saude") in upas]
        campos = request.args.getlist("campos")
        if campos:
            def _tem_campo(pac, campo):
                av = pac.get("avaliacao", {})
                id_ = pac.get("identificacao", {})
                return (
                    (av.get(campo) is not None and av.get(campo) != "")
                    or (id_.get(campo) is not None and id_.get(campo) != "")
                    or (pac.get(campo) is not None and pac.get(campo) != "")
                )

            pacientes = [p for p in pacientes if all(_tem_campo(p, c) for c in campos)]
        unidade_saude = request.args.get("unidade_saude")
        if unidade_saude and not upas and not campos:
            pacientes = db.buscar_pacientes({"unidade_saude": unidade_saude})
        filtros = {
            "inicio_pre_natal": request.args.get("inicio_pre_natal"),
            "plano_parto": request.args.get("plano_parto"),
            "participou_grupos": request.args.get("participou_grupos"),
            "possui_bolsa_familia": request.args.get("possui_bolsa_familia"),
            "tem_vacina_covid": request.args.get("tem_vacina_covid"),
            "vacinas": request.args.get("vacinas"),
        }
        pacientes = aplicar_filtros_exportacao(pacientes, filtros)
        colunas_personalizadas = request.args.getlist("colunas") or None
        if formato == "excel":
            return _exportar_excel(pacientes, colunas_personalizadas)
        if formato == "word":
            return _exportar_word(pacientes, colunas_personalizadas)
        if formato == "txt":
            return _exportar_txt(pacientes, colunas_personalizadas)
        return jsonify({"success": False, "message": "Formato não suportado"}), 400
    except Exception as e:
        return jsonify({"success": False, "message": f"Erro ao exportar: {str(e)}"}), 500
